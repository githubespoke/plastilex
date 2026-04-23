# -*- coding: utf-8 -*-
"""
Generate PlastiLex_Blueprint_ResearchNavigator.docx
Matches structure/styling of existing Option 1 and Option 2 blueprints.
"""

from docx import Document
from docx.shared import Pt


def add_title_block(doc):
    p = doc.add_paragraph()
    run = p.add_run('PlastiLex: Project Blueprint')
    run.font.size = Pt(28)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        'PlastiLex Research Navigator\n'
        'From Search Tool to Research Partner: '
        'AI-Powered Collaborator Discovery for the Chinese Microplastics Literature'
    )
    run2.font.size = Pt(18)
    p2.paragraph_format.space_after = Pt(8)

    p3 = doc.add_paragraph()
    r = p3.add_run('Prepared by: ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = p3.add_run('Spike Incubator')
    r2.font.size = Pt(10)
    p3.add_run('   ')
    r3 = p3.add_run('Student: ')
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p3.add_run('Nyssa Kansal')
    r4.font.size = Pt(10)
    p3.add_run('   ')
    r5 = p3.add_run('Date: ')
    r5.bold = True
    r5.font.size = Pt(10)
    r6 = p3.add_run('April 2026')
    r6.font.size = Pt(10)
    p3.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text):
    doc.add_paragraph(text)


def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def numbered(doc, text):
    doc.add_paragraph(text, style='List Number')


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()


def build():
    doc = Document()
    add_title_block(doc)

    # 1. WHAT IS IT
    h1(doc, 'What Is the PlastiLex Research Navigator?')

    para(doc,
        'NIH and NSF reviewers increasingly flag grant proposals that miss substantial '
        'international literature. In microplastics research, that means Chinese-language '
        'studies -- thousands of them published every year in databases that English-speaking '
        'scientists cannot read, navigate, or cite. The work exists. Most researchers do not '
        'know it does.'
    )
    para(doc,
        'PlastiLex Research Navigator is a web application built to close that gap. '
        'It gives any researcher a semantically searchable index of Chinese microplastics '
        'literature, an AI chat interface that answers questions about that literature, '
        'a collaborator discovery engine that surfaces the Chinese scientists most relevant '
        'to their work -- and a Discover mode that surfaces research gaps they did not know '
        'existed. The result: a stronger, more credible literature review and a grant '
        'proposal that reviewers cannot flag for missing international coverage.'
    )
    para(doc, 'The central shift from any previous PlastiLex concept:')
    para(doc,
        'From: "Here are the gaps in Chinese-English microplastics research coverage."\n'
        'To: "Here are the research areas where Chinese literature can strengthen your '
        'next grant proposal -- and the researchers you should know."'
    )
    para(doc,
        'The gap analysis data is not the product. It is the fuel that makes the product '
        'useful. A researcher arrives -- with or without a specific topic -- and leaves '
        'with: (a) a map of relevant Chinese research they did not know existed, '
        '(b) a shortlist of Chinese scientists working in that space, and '
        '(c) the tools to reach out and cite.'
    )

    # 2. WHY NYSSA
    h1(doc, 'Why Nyssa?')

    para(doc,
        'This project requires a combination of capabilities that almost never exists '
        'in a high school student:'
    )
    bullet(doc,
        'Mandarin reading proficiency at academic level. Six years of Mandarin study at '
        'Harker. The only non-native speaker in her Advanced Mandarin class. Can read, '
        'interpret, and classify Chinese scientific abstracts without machine assistance.'
    )
    bullet(doc,
        "Domain knowledge in microplastics science. Nyssa spent a year in Stanford's "
        'Dirzo Lab testing microplastic effects on organisms. She understands the science '
        'deeply enough to evaluate what a Chinese paper is actually studying -- not just '
        'what its title says.'
    )
    bullet(doc,
        'Builder instinct and AI-development literacy. The same instinct that drove '
        'MycoMap: she identifies a structural gap in how researchers work and builds a '
        'tool to close it. She can use Cursor, Claude Code, and Replit AI to build and '
        'deploy production software.'
    )
    bullet(doc,
        'The NSLI-Y Taiwan connection. If Nyssa participates in the NSLI-Y Mandarin '
        'immersion program in Taiwan, she will have academic-year proximity to '
        'Chinese-language scientific culture while building a tool designed to make '
        'that culture accessible to the world.'
    )
    para(doc,
        "She is not summarizing other people's summaries. She is doing primary "
        'classification of primary sources in their original language -- and then '
        'building an AI system that makes that corpus interactive.'
    )

    # 3. HOW IT WORKS
    h1(doc, 'How It Works')
    h2(doc, 'Two Ways to Use It')

    h3(doc, 'Session A: Directed Search -- For researchers who arrive with a topic (15 minutes)')

    para(doc,
        'Dr. Sarah Chen, postdoc at UC San Diego, is writing a grant on freshwater '
        'microplastic sediment accumulation. She needs to demonstrate awareness of the '
        'full international literature.'
    )
    numbered(doc,
        'She types: "What does Chinese research say about microplastic accumulation '
        'in river sediment?"'
    )
    numbered(doc,
        'The chatbot responds: "There are 34 Chinese-language studies on river sediment '
        'microplastics in our database (2016-2024), concentrated in the Yangtze and Pearl '
        'River basins. This area has roughly 3x the publication density in Chinese vs. '
        'English literature. See the papers, or find researchers active in this area?"'
    )
    numbered(doc,
        'Three author profiles appear: publication records, institutional affiliations, '
        'ORCID links, and international co-authorship signals.'
    )
    numbered(doc,
        'She clicks a profile, sees the paper list, clicks "Draft outreach email." '
        'The tool generates a professional English introduction email.'
    )
    numbered(doc, 'She customizes two sentences and sends it. Total time: under 15 minutes.')

    h3(doc, 'Session B: Discovery Mode -- For researchers who do not know what to ask (10 minutes)')

    para(doc,
        'Dr. James Park, assistant professor at MIT, is curious about Chinese '
        'environmental science but has no specific topic in mind.'
    )
    numbered(doc,
        'He lands on PlastiLex. The homepage IS the Discover view -- no clicks required. '
        'Five Featured Gap cards are visible immediately.'
    )
    numbered(doc,
        'He reads: "Soil microplastics in agricultural land -- 134 Chinese papers, '
        '23 English equivalents. Gap Score 5.8. Chinese researchers are investigating '
        'how plastic mulch film degradation affects crop soil quality at rates far '
        'exceeding Western study."'
    )
    numbered(doc,
        'He clicks the card. The research browser opens filtered to that topic: '
        '134 papers, keyword clusters, and 12 active Chinese researchers.'
    )
    numbered(doc,
        "The chatbot panel pre-loads with context: \"You're exploring soil microplastics "
        'in agricultural land. Which of these researchers have published in English?" '
        'He asks the follow-up.'
    )
    numbered(doc,
        'He finds two bridge researchers with dual-language publication records and '
        'drafts outreach emails to both. Total time: under 10 minutes. From zero '
        'knowledge to researcher contact.'
    )

    h2(doc, 'What Each Featured Gap Card Shows')

    para(doc, 'Every Featured Gap card on the homepage includes:')
    bullet(doc, 'Topic name (e.g., "Soil microplastics in agricultural land")')
    bullet(doc, 'Gap Score (e.g., 5.8 -- ratio of Chinese to English publication count)')
    bullet(doc, 'One-sentence plain-English summary of why this gap matters to researchers')
    bullet(doc, 'Paper count: Chinese vs. English (e.g., "134 Chinese / 23 English")')
    bullet(doc, 'Top 2-3 Chinese researchers active in that area, with institution')
    bullet(doc,
        'Click action: opens research browser filtered to that topic with chatbot '
        'pre-loaded'
    )

    h2(doc, 'Legal Compliance')

    para(doc,
        "PlastiLex operates entirely within the legal constraints identified in the "
        "project's prior research phase:"
    )
    bullet(doc,
        'Metadata-only corpus. No full text, no translations of original content. '
        'Every paper record links to the original source in CNKI or Wanfang -- the '
        'researcher clicks through to the publisher.'
    )
    bullet(doc,
        'Student per-paper annotations are original composition -- 3-5 sentences '
        'written by Nyssa in her own words about what a paper studies, its method, '
        'and its key finding. These are not reproductions of abstract text. They are '
        'the primary RAG (retrieval-augmented generation) source for chatbot answers.'
    )
    bullet(doc,
        'Bibliographic metadata sent to the OpenAI embeddings API is treated as '
        'bibliographic fact -- the same legal basis as indexing citations. No '
        'licensed text is transmitted.'
    )

    # 4. TECHNICAL SPECIFICATION
    h1(doc, 'Technical Specification')

    h2(doc, 'System Architecture')

    para(doc, 'The application has three layers:')
    bullet(doc,
        'Frontend: Browser application with Discover homepage (5 Featured Gap cards), '
        'Research Browser (paper search, author profiles, gap view), and Chat Interface '
        '(right panel, synced with browser navigation).'
    )
    bullet(doc,
        'Backend: FastAPI application deployed on Hugging Face Spaces. Five API endpoints: '
        '/search (semantic paper search), /authors (collaborator matching), '
        '/gaps (pre-computed gap analysis), /gaps/featured (top 5 gaps for Discover '
        'homepage), /chat (RAG chatbot).'
    )
    bullet(doc,
        'Data Layer: Static JSON and SQLite files -- paper_metadata.json '
        '(~800-1,200 records), paper_annotations.json (student-authored summaries), '
        'author_profiles.sqlite, gap_analysis.json, embeddings_cache.json '
        '(pre-computed via OpenAI API).'
    )

    h2(doc, 'AI Components')

    add_table(doc,
        ['Capability', 'Implementation', 'Notes'],
        [
            ('Semantic search',
             'OpenAI embeddings API (text-embedding-3-small)',
             'One-time cost under $0.02 for 1,000 papers; multilingual-capable'),
            ('RAG chatbot',
             'Claude or OpenAI chat API with strict retrieval grounding',
             'Answers grounded in student annotations; shows what it cannot answer'),
            ('Collaborator matching',
             'Cosine similarity on author profile embeddings',
             'Same OpenAI embeddings; no additional tooling'),
            ('Email drafting',
             'LLM prompt on collaborator match',
             'Simple addition; student edits before sending'),
            ('Gap scoring',
             'Math on paper counts',
             'No AI required; deterministic ratio calculation'),
        ]
    )

    h2(doc, 'Tech Stack')

    bullet(doc, 'Frontend: Next.js + Tailwind CSS, or plain HTML/JS for simplicity')
    bullet(doc, 'Backend: FastAPI (Python) on Hugging Face Spaces (free tier, purpose-built for AI demos)')
    bullet(doc, 'Vector index: FAISS (pre-computed, shipped with the app)')
    bullet(doc, 'Embeddings: OpenAI text-embedding-3-small API')
    bullet(doc, 'Chat API: Claude or OpenAI (student choice based on access)')
    bullet(doc, 'Author data: OpenAlex API (free, no auth) + ORCID API (free public API)')
    bullet(doc, 'Data storage: JSON + SQLite -- no database server required')
    bullet(doc, 'Publication: Zenodo DOI for corpus and methodology report')
    bullet(doc, 'Budget: Under $30 total (domain ~$12, OpenAI embeddings ~$0.02, everything else free tier)')

    h2(doc, 'What the Chatbot Can and Cannot Do')

    para(doc, 'Can do (shown in UI):')
    bullet(doc, 'Answer topic queries grounded in the metadata and student annotations')
    bullet(doc, 'Navigate the paper browser (filtering, sorting, narrowing)')
    bullet(doc, 'Surface gap analysis for any queried topic')
    bullet(doc, 'Find researcher matches and explain why they match')
    bullet(doc, 'Draft professional outreach emails in English')
    bullet(doc, 'Pre-load context when a user clicks a Featured Gap card')

    para(doc, 'Cannot do (shown explicitly in UI):')
    bullet(doc, 'Summarize paper full text (legal restriction -- links to originals only)')
    bullet(doc, 'Perform real-time CNKI search (static index, updated periodically)')
    bullet(doc, 'Guarantee accuracy of contact information')

    # 5. EVIDENCE-STACKING STRATEGY
    h1(doc, 'Evidence-Stacking Strategy')
    h2(doc, 'Spike Evidence Pyramid -- Applied to PlastiLex Research Navigator')

    para(doc,
        'Level 1 -- Participation: Built a bilingual AI research tool. Baseline.'
    )
    para(doc,
        'Level 2 -- Achievement/Validation: Corpus published on Zenodo with a permanent DOI. '
        'Methodology report published alongside dataset. Live application at a public URL. '
        'Any researcher worldwide can use, download, or cite the work independently of Nyssa.'
    )
    para(doc,
        'Level 3 -- Character: The project demonstrates cross-cultural scientific '
        'bridge-building as a functional capability. The Discover mode operationalizes it: '
        'the tool does not wait for researchers to know what to ask. It asks better '
        'questions on their behalf.'
    )
    para(doc,
        'Level 4 -- Leadership: The collaborator matching methodology and Gap Score formula '
        'set a reproducible standard. Other students or researchers can extend the corpus '
        'to Spanish (SciELO), Japanese (J-STAGE), or Russian (CyberLeninka) using the '
        'same normalized schema.'
    )
    para(doc,
        'Level 5 -- Impact: Primary evidence: any citation of the PlastiLex corpus by an '
        'external researcher, or any research collaboration that began through the tool. '
        'Secondary: dataset download counts, chatbot session data, researcher feedback. '
        'Tertiary: press coverage or academic acknowledgment of the tool.'
    )

    h2(doc, 'The Recommendation Letter Test')

    para(doc,
        'A Stanford professor can write, truthfully and specifically: "She identified a '
        'structural gap in how the global microplastics research community communicates '
        'across language lines, built an AI system that closes it for any researcher '
        'worldwide, and the tool has been used by researchers at [institutions] to '
        'identify Chinese collaborators they would not otherwise have found."'
    )
    para(doc,
        'That sentence is impossible to manufacture. It describes a verifiable, unique '
        'contribution that no other applicant can claim. The collaborator matching angle '
        'is the strongest element -- it demonstrates systems-level thinking about human '
        'connection, not just data access. The Discover mode adds a second narrative '
        'layer: the tool does not just answer questions, it asks better ones.'
    )

    h2(doc, 'College Application Strength')

    add_table(doc,
        ['Dimension', 'Gap Dashboard (Option 1)', 'Research Navigator (This Project)'],
        [
            ('What student built',
             'Data visualization',
             'Product researchers can actively use'),
            ('Technical skills',
             'Data analysis, web dev',
             'AI/ML, RAG, OpenAI API, vector search, collaborator matching'),
            ('Real-world impact',
             'Shows a problem',
             'Reduces a documented research collaboration barrier'),
            ('Novelty',
             'Bibliometrics tools exist',
             'No free tool cross-matches Chinese/English microplastics researchers'),
            ('Essay hook',
             '"I catalogued a gap"',
             '"I built an AI tool that connects American scientists to Chinese '
             'research partners they did not know existed"'),
            ('Discovery angle',
             'N/A',
             'The tool surfaces gaps you did not know existed -- '
             'you do not need to know what to ask'),
        ]
    )

    # 6. CONTENT STRATEGY
    h1(doc, 'Content Strategy')

    h2(doc, 'Data Sourcing Workflow')

    para(doc, 'Databases: CNKI (China National Knowledge Infrastructure), Wanfang Data')
    para(doc, 'Taxonomy construction (prior to database search):')
    bullet(doc, 'Read 10-15 English review articles to extract recurring sub-topic categories')
    bullet(doc, 'Define 15-25 formal taxonomy tags, each with a one-sentence definition')
    bullet(doc, 'Write a 1-page justification document for taxonomy choices')

    para(doc, 'Per-record workflow (estimated 5-10 minutes per entry):')
    numbered(doc, 'Search CNKI/Wanfang for the taxonomy tag in Chinese')
    numbered(doc,
        'Record the bibliographic metadata (title, authors, journal, year, DOI, '
        'institution, province)'
    )
    numbered(doc, 'Assign taxonomy tags from the formal schema')
    numbered(doc,
        'Flag English analog status (does a comparable English study exist?)'
    )
    numbered(doc,
        "Write 3-5 sentence annotation in Nyssa's own words: what the paper studies, "
        'method used, key finding, geographic context'
    )

    para(doc,
        'V1 target: 800-1,200 records. The per-paper annotations are the most '
        'time-intensive component -- and the most valuable. They are what makes chatbot '
        'answers materially useful, and they are original scholarly composition that '
        'cannot be replicated by a tool.'
    )

    h2(doc, 'Gap Analysis and Featured Gaps')

    para(doc,
        'After the corpus is built, Nyssa computes Gap Scores across all taxonomy tags '
        '(Chinese record count / English-language comparison count) and selects the five '
        'highest-scoring, most research-relevant gaps for the Discover homepage.'
    )
    para(doc, 'Each Featured Gap requires:')
    bullet(doc,
        'One-sentence plain-language summary of why this gap matters to researchers'
    )
    bullet(doc, 'Paper count: Chinese vs. English')
    bullet(doc,
        'Top 2-3 Chinese researchers most active in that sub-topic '
        '(names + institutions + ORCID)'
    )
    bullet(doc, 'Gap Score')

    para(doc,
        'The five Featured Gaps are curated, not auto-generated. Nyssa selects them for '
        'research significance, not just numeric score. This curation is itself part of '
        'the original scholarly contribution.'
    )

    h2(doc, 'Author Profiles and Collaborator Matching')

    para(doc, 'Author profiles are built from three sources:')
    bullet(doc, 'CNKI/Wanfang export: Chinese-side publication records')
    bullet(doc,
        'OpenAlex API: English-side publication records, citations, international '
        'co-authors'
    )
    bullet(doc,
        'ORCID API: Affiliation links, disambiguation keys, existing international '
        'collaborations'
    )

    para(doc,
        'Bridge researchers -- those appearing in both Chinese and English corpora -- '
        'are highest-priority collaborator matches and are flagged prominently.'
    )

    h2(doc, 'V1.5 Roadmap (Post-Launch Extensions)')

    para(doc, 'If Nyssa completes V1 with time remaining:')
    bullet(doc,
        'Keyword co-occurrence map: Force-directed network graph showing keyword '
        'clusters in Chinese vs. English research. Estimated 10-15 additional hours. '
        'Deferred from V1 by CEO decision -- Featured Gaps deliver 80% of discovery '
        'value at 20% of the effort.'
    )
    bullet(doc,
        'Additional language corpora: Spanish/Portuguese (SciELO), Japanese (J-STAGE) '
        'following the same normalized schema. Architecture supports multilingual '
        'extension from day one.'
    )
    bullet(doc,
        'User accounts and saved searches: Allow researchers to bookmark gaps and '
        'researchers.'
    )

    # 7. EXECUTION TIMELINE
    h1(doc, 'Execution Timeline')

    h2(doc, 'Week 1: Corpus Design and Data Acquisition')
    bullet(doc, 'Define taxonomy; CNKI/Wanfang export (3 sessions x ~400 records); OpenAlex API pull')
    bullet(doc, 'Write data acquisition protocol')
    bullet(doc, 'Deliverables: chinese-corpus-raw.csv, english-corpus-raw.csv, acquisition-protocol.md')

    h2(doc, 'Week 2: Data Cleaning and Author Profile Construction')
    bullet(doc, 'Normalize author names, affiliations, keywords')
    bullet(doc, 'Build author_profiles.sqlite with ORCID enrichment; flag bilingual researchers')
    bullet(doc, 'Deliverables: paper-metadata-clean.csv, author-profiles.sqlite')

    h2(doc, 'Week 3: Per-Paper Annotations (CTO Requirement for RAG Quality)')
    bullet(doc,
        'Student writes 3-5 sentence annotation per paper: what it studies, method '
        'used, key finding, geographic context. These are the primary RAG retrieval '
        'documents.'
    )
    bullet(doc, 'Deliverable: paper-annotations.json (~800-1,200 entries)')

    h2(doc, 'Week 4: Gap Analysis Computation')
    bullet(doc, 'Compute Gap Scores per taxonomy tag')
    bullet(doc,
        'Write 5 Featured Gap narratives (one-sentence summaries + researcher shortlists)'
    )
    bullet(doc, 'Identify top 2-3 researchers per featured gap')
    bullet(doc, 'Deliverables: gap-analysis.json, featured-gaps.json')

    h2(doc, 'Week 5: Embeddings and Vector Index')
    bullet(doc, 'Call OpenAI text-embedding-3-small on all paper titles + annotations')
    bullet(doc, 'Build FAISS index for papers and author profiles')
    bullet(doc, 'Test semantic search with 10 sample queries')
    bullet(doc, 'Deliverables: embeddings_cache.json, faiss-paper.index, faiss-author.index')

    h2(doc, 'Week 6: FastAPI Backend')
    bullet(doc, 'Build five endpoints: /search, /authors, /gaps, /gaps/featured, /chat')
    bullet(doc, 'Deploy to Hugging Face Spaces')
    bullet(doc, 'Write automated tests')
    bullet(doc, 'Deliverable: Live HF Spaces URL, test suite')

    h2(doc, 'Week 7: RAG Chatbot')
    bullet(doc, 'System prompt with strict retrieval grounding')
    bullet(doc, 'Integrate chat API; context retrieval pipeline; email draft generation')
    bullet(doc, 'Test 20 canonical queries')
    bullet(doc,
        'Test Discover-mode contextual prompts (chatbot pre-loaded with gap context '
        'when user clicks a Featured Gap)'
    )
    bullet(doc, 'Deliverable: Working chat endpoint, test conversation log')

    h2(doc, 'Week 8: Frontend -- Discover Homepage + Research Browser')
    bullet(doc,
        'Featured Gaps homepage with 5 gap cards (topic, score, summary, paper counts, '
        'top researchers)'
    )
    bullet(doc, 'Gap card click -> filtered research browser')
    bullet(doc,
        'Paper search/filter interface; gap analysis view; author profile pages with '
        'collaborator matching'
    )
    bullet(doc, 'Deliverable: Discover homepage + browser interface connected to backend')

    h2(doc, 'Week 9: Frontend -- Chat Integration')
    bullet(doc, 'Chat panel with bidirectional sync to main panel')
    bullet(doc, 'Contextual chat prompt on Discover homepage')
    bullet(doc, 'Mobile-responsive layout')
    bullet(doc,
        'Deliverable: Full UI working end-to-end; both Session A and Session B flows '
        'verified'
    )

    h2(doc, 'Week 10: User Testing and Polish')
    bullet(doc, 'Share with 1-2 actual researchers')
    bullet(doc,
        'Test the Session B (Discovery) flow specifically: does a researcher with no '
        'topic find something interesting within 5 minutes?'
    )
    bullet(doc, 'Fix top-3 usability issues; write Zenodo methodology report')
    bullet(doc, 'Deliverables: User test notes, methodology draft')

    h2(doc, 'Week 11: Publication and Application Documentation')
    bullet(doc, 'Zenodo record, GitHub README, Loom demo video, college application writeup')
    bullet(doc,
        'Demo video leads with the Discovery flow -- it is the most visually compelling '
        'for admissions readers who check the project URL'
    )
    bullet(doc, 'Deliverables: Public Zenodo DOI, GitHub repo, demo video')

    # 8. WHAT SUCCESS LOOKS LIKE
    h1(doc, 'What Success Looks Like')

    para(doc, 'By the end of 11 weeks, Nyssa will have:')
    numbered(doc,
        'A live, publicly accessible PlastiLex Research Navigator with 800-1,200 curated '
        'metadata records and per-paper annotations'
    )
    numbered(doc,
        'An AI chatbot that answers real researcher questions about the Chinese '
        'microplastics literature -- grounded in original annotations, not hallucinated'
    )
    numbered(doc,
        'A collaborator discovery engine that matches researchers to Chinese scientists '
        'by topic and flags bridge researchers with dual-language publication records'
    )
    numbered(doc,
        'A Discover mode homepage with 5 Featured Gap cards -- the most compelling '
        'demonstration that the tool asks better questions than users know to ask'
    )
    numbered(doc,
        'A Zenodo-published open corpus with permanent DOI -- independently citable '
        'by any researcher worldwide'
    )
    numbered(doc,
        'A formal methodology report and gap summary brief published on Zenodo'
    )
    numbered(doc,
        'Documented usage data: Zenodo download counts, chatbot session metrics, '
        'site analytics'
    )
    numbered(doc,
        'An authentic, verifiable narrative: she built an AI tool that connects '
        'English-speaking scientists to Chinese research partners they did not know '
        'existed -- and the tool asks better questions than they knew to ask'
    )
    numbered(doc,
        'A recommendation letter pathway: any faculty member or researcher who uses '
        'the tool to find a collaborator can write specifically and verifiably about '
        'a real contribution'
    )

    # 9. RISKS AND MITIGATIONS
    h1(doc, 'Risks and Mitigations')

    add_table(doc,
        ['Risk', 'Probability', 'Impact', 'Mitigation'],
        [
            ('Annotation workload (800-1,200 x 3-5 sentences)',
             'High', 'Medium',
             'Timebox Week 3; reduce corpus to 500 records if needed. '
             'Retrieval still works at 500.'),
            ('RAG quality',
             'Medium (CTO: viable with good annotations)', 'High',
             'Start retrieval-only; add generation after retrieval is validated'),
            ('CNKI export friction',
             'High', 'Low',
             '3 sessions planned; 400 records minimum is sufficient for V1'),
            ('Author disambiguation',
             'Medium', 'Medium',
             'ORCID as primary key; confidence scores displayed on all collaborator '
             'matches'),
            ('Scope creep',
             'High', 'High',
             'Hard MVP: Discover homepage + search + gap view + chatbot + collaborator '
             'matching. Cut corpus size or Week 10 polish before cutting features'),
            ('OpenAlex Chinese coverage gap',
             'High', 'Low',
             'Expected; ORCID supplements for disambiguation. Flagged in UI.'),
            ('Featured Gap curation quality',
             'Low', 'Medium',
             'Student selects top 5 from gap scores; board review before launch'),
            ('Keyword co-occurrence map complexity',
             'High', 'Low',
             'Deferred to V1.5 by design. Not on critical path.'),
        ]
    )

    # FOOTER
    p_footer = doc.add_paragraph(
        'Blueprint prepared by Spike Incubator. '
        'For questions, contact Hari Iyer or Ro Arora at whatsyourspike.com.'
    )
    p_footer.paragraph_format.space_before = Pt(24)

    out = 'mockup/PlastiLex_Blueprint_ResearchNavigator.docx'
    doc.save(out)
    print('Saved: ' + out)


if __name__ == '__main__':
    build()
