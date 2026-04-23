"""
Generate PlastiLex_Permissions_Summary.docx
Parent-facing document explaining why seeking translation permissions is not viable
and why the Research Navigator pivot is the right move.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_heading_style(para, level=1):
    """Apply heading formatting."""
    run = para.runs[0] if para.runs else para.add_run(para.text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)  # dark navy
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x3E)  # dark green
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_callout_box(doc, text, bg_color="F0F4F8"):
    """Add a styled callout box paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x2E)
    return p


doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# ── Title block ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title.add_run("PlastiLex: Project Direction Update")
t_run.font.size = Pt(22)
t_run.font.bold = True
t_run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = sub.add_run("Why We Pivoted — and Why the New Direction Is Stronger")
s_run.font.size = Pt(13)
s_run.font.italic = True
s_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
m_run = meta.add_run("Prepared for: Parent / College Application Reviewer  |  PlastiLex Project  |  April 2026")
m_run.font.size = Pt(9)
m_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_horizontal_rule(doc)
doc.add_paragraph()  # spacer


# ── Section 1 ─────────────────────────────────────────────────────────────────
h1 = doc.add_paragraph()
r = h1.add_run("1.  What the Original Plan Was")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "PlastiLex began with an ambitious goal: take Chinese-language academic research on microplastics — "
    "some of the most extensive in the world — and make it accessible to English-speaking scientists, "
    "students, and policymakers. The plan was to translate papers sourced from the three dominant Chinese "
    "academic databases (CNKI, Wanfang, and CQVIP) and publish them on Zenodo, an open-access research "
    "repository, and through a companion web application."
)
p.paragraph_format.space_after = Pt(8)

p2 = doc.add_paragraph(
    "The vision was genuine and the problem was real: a significant body of Chinese-language microplastics "
    "research is effectively invisible to Western researchers because of the language barrier. Bridging that "
    "gap would have been a meaningful contribution."
)
p2.paragraph_format.space_after = Pt(10)


# ── Section 2 ─────────────────────────────────────────────────────────────────
h2 = doc.add_paragraph()
r = h2.add_run("2.  What the Research Revealed")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "Before building anything, we did a thorough legal review of each database's terms of service — "
    "exactly the kind of due diligence a responsible project should do. What we found was unambiguous:"
)
p.paragraph_format.space_after = Pt(6)

bullets = [
    ("CNKI", "Grants users only personal, non-transferable rights. Translation and republication are "
     "explicitly prohibited without written authorization from the rights holder."),
    ("Wanfang Data", "Same restriction. Personal use only. Derivative works — including translations — "
     "require separate written permission."),
    ("CQVIP", "Consistent with the others. Unauthorized translation or redistribution is not permitted "
     "under the platform's terms."),
]

for label, text in bullets:
    bp = doc.add_paragraph(style="List Bullet")
    bold_run = bp.add_run(label + ": ")
    bold_run.bold = True
    bold_run.font.size = Pt(11)
    rest = bp.add_run(text)
    rest.font.size = Pt(11)
    bp.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph(
    "Under Chinese Copyright Law (Article 10), translation is an exclusive right held by the copyright "
    "owner. The \"fair use\" provision (Article 24) covers personal study only — it does not extend to "
    "publishing or distributing translated work, even for educational purposes."
)
p3.paragraph_format.space_before = Pt(8)
p3.paragraph_format.space_after = Pt(10)


# ── Section 3 ─────────────────────────────────────────────────────────────────
h3 = doc.add_paragraph()
r = h3.add_run("3.  Why Seeking Permission Is Not a Realistic Path")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "The natural next question is: \"Can we just ask for permission?\" We looked into that carefully. "
    "Here is why that path does not work for a project of this scope and timeline:"
)
p.paragraph_format.space_after = Pt(8)

reasons = [
    ("Three separate gatekeepers, each with different processes.",
     "CNKI, Wanfang, and CQVIP each require independent permission requests directed to different "
     "departments. There is no single \"Chinese academic database\" authority to contact."),
    ("Timelines that do not fit a project timeline.",
     "Each database's response window runs 4–12 weeks — and that is only to receive a decision, not "
     "to complete the paperwork. With three databases, the process could stretch six months or more "
     "before a single paper could be published."),
    ("No guarantee of approval.",
     "Success rates were assessed at 20–50% per database — and that is for straightforward commercial "
     "requests. A student educational project may receive less priority, not more."),
    ("\"Educational use\" is not a shortcut.",
     "This is a common misconception. Chinese copyright law does not provide a blanket educational "
     "exemption for publishing translated works. Permission must still be obtained."),
    ("Real legal risk if permission is not granted.",
     "CNKI has an established track record of enforcing its copyright through litigation. Court "
     "precedents in China have resulted in damages exceeding 700,000 yuan. This is not a theoretical "
     "risk — it is documented."),
]

for i, (heading, explanation) in enumerate(reasons, 1):
    np = doc.add_paragraph()
    np.paragraph_format.left_indent = Inches(0.2)
    np.paragraph_format.space_after = Pt(6)
    num = np.add_run(f"{i}.  ")
    num.bold = True
    num.font.size = Pt(11)
    num.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)
    head_run = np.add_run(heading + "  ")
    head_run.bold = True
    head_run.font.size = Pt(11)
    body_run = np.add_run(explanation)
    body_run.font.size = Pt(11)

doc.add_paragraph()  # spacer


# ── Section 4 ─────────────────────────────────────────────────────────────────
h4 = doc.add_paragraph()
r = h4.add_run("4.  What We Can Do — Legally and Ethically")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "The legal review did not close every door. Several valuable things remain fully permissible:"
)
p.paragraph_format.space_after = Pt(8)

ok_items = [
    "Publishing metadata — titles, authors, journals, publication years, and DOIs",
    "Linking to original papers on each database",
    "Building discovery and analysis tools around the research without reproducing content",
    "Creating original analysis of trends, patterns, and gaps in the research landscape",
    "Producing original commentary, synthesis, and summaries written from scratch",
]

for item in ok_items:
    bp = doc.add_paragraph(style="List Bullet")
    r = bp.add_run(item)
    r.font.size = Pt(11)
    bp.paragraph_format.space_after = Pt(4)

doc.add_paragraph()  # spacer


# ── Section 5 ─────────────────────────────────────────────────────────────────
h5 = doc.add_paragraph()
r = h5.add_run("5.  The Pivot: PlastiLex Research Navigator")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "Rather than stopping at \"we can't translate,\" we asked a better question: "
    "what do English-speaking researchers actually need? "
    "The answer is not translated papers. It is discoverability. "
    "Most researchers do not know what Chinese microplastics research exists, who is doing it, "
    "or how to connect with those teams. The language barrier is real — but the deeper problem is "
    "that there is no bridge at all."
)
p.paragraph_format.space_after = Pt(8)

p2 = doc.add_paragraph(
    "The PlastiLex Research Navigator is that bridge. It is a searchable, structured tool that surfaces "
    "Chinese microplastics research by topic, region, institution, and research gap — letting English-speaking "
    "scientists discover work they could not find before, and identify opportunities for international "
    "collaboration. No translated content. No copyright exposure. Entirely original architecture built "
    "on publicly available metadata and original analysis."
)
p2.paragraph_format.space_after = Pt(8)

add_callout_box(
    doc,
    "\"The Research Navigator does not replace Chinese papers — it makes them findable. That is a more "
    "durable and more useful contribution than a translated archive that could be taken down over a rights dispute.\""
)

doc.add_paragraph()  # spacer


# ── Section 6 ─────────────────────────────────────────────────────────────────
h6 = doc.add_paragraph()
r = h6.add_run("6.  Why This Is Actually Stronger for a College Application")
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)

p = doc.add_paragraph(
    "College admissions readers are looking for evidence of real thinking — not just effort, but "
    "judgment. This project demonstrates several things that are difficult to fake:"
)
p.paragraph_format.space_after = Pt(8)

strengths = [
    ("Identified a genuine problem.",
     "The language barrier in Chinese microplastics research is real, documented, and consequential. "
     "This was not a manufactured school project."),
    ("Investigated it rigorously.",
     "Rather than assuming translation was fine, the project commissioned a legal analysis of three "
     "separate databases' terms of service and traced the relevant provisions of Chinese copyright law. "
     "That is a level of intellectual seriousness that stands out."),
    ("Adapted to what the evidence showed.",
     "When the legal analysis revealed that the original plan was not viable, the project did not "
     "collapse — it pivoted. The new direction was not a fallback; it was a better answer to the "
     "underlying problem."),
    ("Built something original.",
     "The Research Navigator is not a derivative work. It required original architectural design, "
     "original analysis of research gaps, and original thinking about how to serve the scientific "
     "community. That is a more impressive outcome than a translation archive."),
    ("Showed ethical judgment.",
     "The decision to respect copyright — even when a workaround might have been technically possible — "
     "reflects integrity that is worth noting explicitly."),
]

for heading, explanation in strengths:
    bp = doc.add_paragraph(style="List Bullet")
    bold_run = bp.add_run(heading + "  ")
    bold_run.bold = True
    bold_run.font.size = Pt(11)
    rest = bp.add_run(explanation)
    rest.font.size = Pt(11)
    bp.paragraph_format.space_after = Pt(5)

doc.add_paragraph()  # spacer
add_horizontal_rule(doc)
doc.add_paragraph()

# ── Closing ───────────────────────────────────────────────────────────────────
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
c_run = closing.add_run(
    "This project began with ambition, encountered a real obstacle, and came out the other side with "
    "something better. That is not a story of failure. It is a story of how good research actually works."
)
c_run.font.size = Pt(11)
c_run.font.italic = True
c_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "mockup/PlastiLex_Permissions_Summary.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
